import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Explicit file list (only important source files) ─────────────────────────
IMPORTANT_FILES = [
    # Backend
    ("backend", [
        "server.js",
    ]),
    # Frontend
    ("frontend/src", [
        "App.js",
        "index.js",
        "config.js",
    ]),
    ("frontend/src/components", [
        "Dashboard.js",
        "FileUpload.js",
        "InvestmentAdvice.js",
    ]),
    ("frontend/public", [
        "index.html",
    ]),
    # ML Service
    ("ml-service", [
        "main.py",
        "database.py",
        "db_manager.py",
        "investment_advisor.py",
        "data_cleaning_pipeline.py",
        "adaptive_cleaner.py",
        "online_learning.py",
    ]),
    # App services
    ("app/services", [
        "classifier.py",
        "summary.py",
    ]),
    ("app/utils", [
        "text_cleaner.py",
    ]),
    # Training
    ("training", [
        "train_model.py",
        "text_cleaner.py",
        "benchmark_performance.py",
    ]),

]

OUTPUT_FILE = "All_Source_Code.docx"

# Colour palette
CLR_TITLE_BG  = RGBColor(0x1E, 0x1E, 0x2E)
CLR_TITLE_FG  = RGBColor(0xFF, 0xFF, 0xFF)
CLR_FOLDER_BG = RGBColor(0x2D, 0x4A, 0x7A)
CLR_FOLDER_FG = RGBColor(0xFF, 0xFF, 0xFF)
CLR_FILE_BG   = RGBColor(0xE8, 0xF0, 0xFE)
CLR_FILE_FG   = RGBColor(0x1A, 0x1A, 0x2E)
CLR_CODE_BG   = RGBColor(0xF8, 0xF8, 0xF8)
CLR_CODE_FG   = RGBColor(0x1E, 0x1E, 0x1E)
CLR_LINE_NUM  = RGBColor(0x99, 0x99, 0x99)
CLR_BORDER    = RGBColor(0xCC, 0xCC, 0xCC)
# ─────────────────────────────────────────────────────────────────────────────

document = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in document.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_para_shading(para, rgb: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    pPr.append(shd)

def set_para_border_bottom(para, rgb: RGBColor):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_toc(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    for tag, val in [('begin', None), ('instrText', 'TOC \\o "1-3" \\h \\z \\u'), ('separate', None), ('end', None)]:
        if tag == 'instrText':
            el = OxmlElement('w:instrText')
            el.text = val
        else:
            el = OxmlElement('w:fldChar')
            el.set(qn('w:fldCharType'), tag)
        run._r.append(el)
    doc.add_paragraph()

# ── Cover page ────────────────────────────────────────────────────────────────
cover = document.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_shading(cover, CLR_TITLE_BG)
run = cover.add_run('\n\n  Complete Source Code Documentation  \n\n')
run.font.name = 'Calibri'
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = CLR_TITLE_FG

sub = document.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_shading(sub, CLR_TITLE_BG)
r = sub.add_run('Auto-generated — All project source files\n\n')
r.font.name = 'Calibri'
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)

document.add_page_break()

# ── Project File Structure ────────────────────────────────────────────────────
struct_heading = document.add_paragraph()
sh_run = struct_heading.add_run('Project File Structure')
sh_run.font.name = 'Calibri'
sh_run.font.size = Pt(16)
sh_run.font.bold = True
sh_run.font.color.rgb = CLR_FOLDER_BG
set_para_border_bottom(struct_heading, CLR_FOLDER_BG)
struct_heading.paragraph_format.space_after = Pt(6)

SKIP_STRUCT = {
    'node_modules', 'venv', '.git', '__pycache__', '.vscode',
    'build', 'dist', '.next', '.nuxt', 'coverage', 'uploads',
    'ml-service',  # nested duplicate folder inside ml-service
}

def build_tree(folder_path, prefix='', is_root=False):
    try:
        entries = sorted(os.listdir(folder_path))
    except PermissionError:
        return
    # filter out skipped folders and hidden/binary files
    entries = [e for e in entries if e not in SKIP_STRUCT and not e.startswith('.')]
    for i, entry in enumerate(entries):
        full = os.path.join(folder_path, entry)
        connector = '└── ' if i == len(entries) - 1 else '├── '
        extension = '    ' if i == len(entries) - 1 else '│   '
        icon = '📁 ' if os.path.isdir(full) else '📄 '
        line = prefix + connector + icon + entry
        p = document.add_paragraph()
        set_para_shading(p, CLR_CODE_BG)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        r = p.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        r.font.color.rgb = CLR_CODE_FG
        if os.path.isdir(full) and entry not in SKIP_STRUCT:
            build_tree(full, prefix + extension)

# root entry
root_p = document.add_paragraph()
set_para_shading(root_p, CLR_CODE_BG)
root_p.paragraph_format.space_before = Pt(0)
root_p.paragraph_format.space_after  = Pt(0)
root_p.paragraph_format.left_indent  = Cm(0.5)
rr = root_p.add_run('📁 ' + os.path.basename(os.getcwd()) + '/')
rr.font.name = 'Consolas'
rr.font.size = Pt(9)
rr.font.bold = True
rr.font.color.rgb = CLR_FOLDER_BG

build_tree(os.getcwd())

document.add_page_break()

# ── TOC ───────────────────────────────────────────────────────────────────────
toc_heading = document.add_paragraph()
toc_run = toc_heading.add_run('Table of Contents')
toc_run.font.name = 'Calibri'
toc_run.font.size = Pt(16)
toc_run.font.bold = True
toc_run.font.color.rgb = CLR_FOLDER_BG
set_para_border_bottom(toc_heading, CLR_FOLDER_BG)

add_toc(document)
document.add_page_break()

# ── Write sections ────────────────────────────────────────────────────────────
base = os.getcwd()

for folder, files in IMPORTANT_FILES:
    folder_path = os.path.join(base, folder)
    folder_label = folder if folder != '.' else 'root'

    # Folder banner
    folder_para = document.add_paragraph()
    set_para_shading(folder_para, CLR_FOLDER_BG)
    fr = folder_para.add_run(f'  📁  {folder_label}')
    fr.font.name = 'Calibri'
    fr.font.size = Pt(14)
    fr.font.bold = True
    fr.font.color.rgb = CLR_FOLDER_FG
    folder_para.paragraph_format.space_before = Pt(12)
    folder_para.paragraph_format.space_after  = Pt(2)

    for filename in files:
        full_path = os.path.join(folder_path, filename)
        if not os.path.isfile(full_path):
            continue

        # File header
        file_para = document.add_paragraph()
        set_para_shading(file_para, CLR_FILE_BG)
        set_para_border_bottom(file_para, CLR_BORDER)
        file_run = file_para.add_run(f'    📄  {filename}')
        file_run.font.name = 'Calibri'
        file_run.font.size = Pt(12)
        file_run.font.bold = True
        file_run.font.color.rgb = CLR_FILE_FG
        file_para.paragraph_format.space_before = Pt(8)
        file_para.paragraph_format.space_after  = Pt(0)

        # Read source
        try:
            with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                code = f.read().strip()
        except Exception:
            code = '⚠️  Could not read file.'

        if not code:
            code = '# (empty file)'

        # Code block with line numbers
        for i, line in enumerate(code.splitlines(), start=1):
            code_para = document.add_paragraph()
            set_para_shading(code_para, CLR_CODE_BG)
            code_para.paragraph_format.space_before = Pt(0)
            code_para.paragraph_format.space_after  = Pt(0)
            code_para.paragraph_format.left_indent  = Cm(0.3)

            ln_run = code_para.add_run(f'{i:>4}  ')
            ln_run.font.name = 'Cambria'
            ln_run.font.size = Pt(8)
            ln_run.font.color.rgb = CLR_LINE_NUM

            code_run = code_para.add_run(line)
            code_run.font.name = 'Cambria'
            code_run.font.size = Pt(10)
            code_run.font.color.rgb = CLR_CODE_FG

        # Spacer
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after  = Pt(8)

# ── Save ──────────────────────────────────────────────────────────────────────
document.save(OUTPUT_FILE)
print(f'✅ Done! Word file created: {OUTPUT_FILE}')
